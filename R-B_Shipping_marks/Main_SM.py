from os import remove
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docxcompose.composer import Composer
from easygui import multenterbox, fileopenbox


def main():
    input_anzahl, input_pages, input_created_doc_name, input_load_doc_path = user_input()

    make_improved_document(input_load_doc_path, input_anzahl)

    compose_doc(input_created_doc_name, "temporary_copy.docx", input_pages)

    change_style_paragraphs_tables(input_created_doc_name, input_pages, input_created_doc_name)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def make_improved_document(load_file, anzahl):
    filename = load_file

    # if '_MEIPASS2' in os.environ:
    #     filename = os.path.join(os.environ['_MEIPASS2'], filename)

    doc = Document(filename)

    table = doc.tables[0]
    table.cell(4, 1).text = anzahl

    delete_paragraph(doc.paragraphs[1])
    doc.add_paragraph("")

    doc.save("temporary_copy.docx")


def compose_doc(created, copy, pages):
    doc_main = Document()
    doc_main.save("temporary_main.docx")
    master = Document("temporary_main.docx")
    composer = Composer(master)
    doc = Document(copy)
    for i in range(pages):
        composer.append(doc)
    composer.save(created)
    remove("temporary_main.docx")
    remove("temporary_copy.docx")


def change_style_paragraphs_tables(load_doc, pages, saved_doc):
    composed_doc = Document(load_doc)
    style = composed_doc.styles["Normal"]
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    delete_counter = len(composed_doc.paragraphs)  # na jedné kopii je 5 paragr., rozlišit sudý/lichý počet stran
    if delete_counter % 2 == 0:                    # umazat čtyři paragr. za tabulkou pro rozložení 2x na list
        for par in range(1, len(composed_doc.paragraphs), 10):
            for i in range(1, 4):
                delete_paragraph(composed_doc.paragraphs[delete_counter - i])
            delete_counter -= 10
    else:
        for par in range(1, len(composed_doc.paragraphs), 10):
            for i in range(1, 4):
                delete_paragraph(composed_doc.paragraphs[delete_counter - i - 5])
            delete_counter -= 10

    table_count = 1
    for tab in range(pages):
        table2 = composed_doc.tables[tab]
        table2.cell(6, 1).text = f"{table_count}-{pages}"  # úprava buněk s počtem stran
        table_count += 1
        for c in range(len(table2.columns)):  # úprava formátu textu buněk
            for r in range(len(table2.rows)):
                copy = table2.cell(r, c).text
                table2.cell(r, c).style = composed_doc.styles['Normal']
                table2.cell(r, c).text = copy

    for table in composed_doc.tables:  # zarovnání první buňky nastřed adresy
        table.cell(0, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    composed_doc.save(saved_doc)


def user_input():
    msg = "Enter your personal information"
    title = "Credit Card Application"
    fieldNames = ["Anzahl", "Number of pages:", "Name of created doc:"]
    fieldValues = ["5 Stück im Karton / 30  Stück im Pallet"]  # we start with blanks for the values
    fieldValues = multenterbox(msg, title, fieldNames, fieldValues)

    # make sure that none of the fields was left blank
    while 1:
        if fieldValues == None:
            break
        errmsg = ""
        try:
            page = int(fieldValues[1])
        except:
            errmsg = errmsg + "Number of pages must be a number"
        for i in range(len(fieldNames)):
          if fieldValues[i].strip() == "":
            errmsg = errmsg + ('"%s" is a required field.\n\n' % fieldNames[i])
        if errmsg == "":
            break  # no problems found
        fieldValues = multenterbox(errmsg, title, fieldNames, fieldValues)

    fieldValues.append(fileopenbox())
    fieldValues[1] = int(fieldValues[1])
    fieldValues[2] += ".docx"
    print("Reply was:", fieldValues)
    return fieldValues


if __name__ == '__main__':
 main()
